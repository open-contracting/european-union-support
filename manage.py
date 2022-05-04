#!/usr/bin/env python
import csv
import os
import re
from pathlib import Path
from textwrap import dedent
from zipfile import ZipFile

import click
import pandas as pd
from docx import Document

basedir = Path(__file__).resolve().parent
sourcedir = basedir / 'source'
mappingdir = basedir / 'output' / 'mapping'
eformsdir = mappingdir / 'eForms'


def check(actual, expected, noun):
    assert actual == expected, f'expected {expected} {noun}, got {actual}'
    return actual


def excel_files():
    with ZipFile(sourcedir / 'Task 5_Support_Standard Forms-eForms mappings_v.3.zip') as zipfile:
        for name in zipfile.namelist():
            with zipfile.open(name) as fileobj:
                with pd.ExcelFile(fileobj) as xlsx:
                    yield name, xlsx


# https://github.com/pallets/click/issues/486
@click.group(context_settings={'max_content_width': 150})
def cli():
    pass


@cli.command()
@click.argument('sheet')
def find(sheet):
    """
    Find the workbook containing the SHEET.
    """
    for name, xlsx in excel_files():
        if sheet in xlsx.sheet_names:
            click.echo(f'First occurrence: {name}')
            break
    else:
        click.secho(f"Sheet '{sheet}' not found", fg='red')


@cli.command()
def extract_xpath_mapping():
    """
    Extract a mapping between Business Terms and eForms XPaths.

    \b
    Create or update output/mapping/eForms/bt-xpath-mapping.csv from source/XPATHs provisional release v. 1.0.docx
    """

    def text(row):
        # Newlines occur in all columns except the first, e.g.:
        # /*/cac:ContractingParty/cac:ContractingRepresentationType/cbc:RepresentationTypeCode
        # /*/cac:ProcurementProjectLot/cac:TenderingTerms/cac:TenderRecipientParty/cbc:EndpointID
        #
        # Leading or trailing whitespace occur in the first and third columns, e.g.:
        # /*/cac:ProcurementProjectLot/cac:ProcurementProject/cbc:EstimatedOverallContractQuantity
        # /*/ext:UBLExtensions/ext:UBLExtension/ext:ExtensionContent/efext:EformsExtension/efac:Change/efbc:C...
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        # XXX: Correct a typo.
        cells[0] = cells[0].replace('[@n listName=', '[@listName=')
        return cells

    docx = Document(sourcedir / 'XPATHs provisional release v. 1.0.docx')
    columns = text(docx.tables[0].rows[0])

    check(len(docx.tables), 1, 'table')
    check(columns, ['XPATH', 'BT ID', 'BT Name', 'Additional information'], 'headers')

    data = []
    for row in docx.tables[0].rows[1:]:
        cells = text(row)
        # Skip subheadings.
        if not cells[0].startswith('/'):
            continue
        # XXX: Correct a typo (n-dash instead of empty).
        if cells[1] == '–':
            cells[1] = ''
        data.append(cells)

    # Note: 4 XPaths map to multiple BTs. (Otherwise, BT:XPath is 1:n.)
    #
    # - /*/cac:TenderingTerms/cac:ProcurementLegislationDocumentReference/cbc:DocumentDescription BT-01 BT-09
    # - /*/cac:TenderingTerms/cac:ProcurementLegislationDocumentReference/cbc:ID BT-01 BT-09
    # - /*/cac:ProcurementProjectLot/cac:ProcurementProject/cac:ProcurementAdditionalType/ext:UBLExtensions/ext:UBLExtension/ext:ExtensionContent/cbc:Description BT-755 BT-777 # noqa: E501
    # - /ContractAwardNotice/cac:TenderResult/cac:AwardedTenderedProject/cac:ProcurementProjectLot/cac:TenderingProcess/cac:FrameworkAgreement/cbc:MaximumValueAmount BT-156 BT-709 # noqa: E501

    pd.DataFrame(data, columns=columns).drop_duplicates().to_csv(eformsdir / 'bt-xpath-mapping.csv', index=False)


@cli.command()
def extract_indices_mapping():
    """
    Extract a mapping between Business Terms and form indices.

    \b
    Create or update output/mapping/eForms/bt-indices-mapping.csv from
    source/Task 5_Support_Standard Forms-eForms mappings_v.3.zip

    The source files don't cover forms 9, 14 and E*: https://docs.ted.europa.eu/home/eforms/FAQ/index.html
    """
    ignore_sheets = {
        'Legend',  # A legend (same for all workbooks).
        'S.F. vs eForms mapping list ',  # A mapping from F## to eForms## (same for all workbooks).
        'Annex table 2',  # A copy of the 2019 regulation's annex.
        'Legend Annex table 2',
    }
    # We only need the 2019-to-2015 direction.
    ignore_sheet_regex = [re.compile(pattern) for pattern in (
        r'^F\d\d vs eN\d\d( \(2\))?$',
        r'^SF\d\d vs eForm ?\d\d(,\d\d)*$'
    )]
    keep_sheet_regex = re.compile(r'^(?:eForm|eN) ?(\d\d?(?:,\d\d)*) vs S?F(\d\d) ?$')

    replace_newlines = re.compile(r'\n(?=\(|(2009|Title III|and|requirements|subcontractor|will)\b)')

    # XXX: Correct the source file's incorrect mapping between 2019 BTs and 2015 indices.
    remove_standard = {
        # Additional Information is top-level like VI.3, not lot-specific.
        'BT-300': 'II.2.14',
        # Tender Value Lowest and Tender Value Highest are lot-specific like V.2.4.3 and V.2.4.4, not top-level.
        'BT-710': 'II.1.7.2',
        'BT-711': 'II.1.7.3',
        # Submission Language has nothing to do with this (/CONTRACTING_BODY/ADDRESS_PARTICIPATION).
        'BT-97': 'I.3.4.1.2',
    }
    remove_concession = {
        **remove_standard,
        # Estimated Value is top-level like II.1.5.1, not lot-specific.
        'BT-27': 'V.2.4.1',
        # Concession Value Description is lot-specific like V.2.4.6, not top-level.
        'BT-163': 'II.1.5.3',
    }
    ignore = {
        # Direct Award Justification Previous Procedure Identifier matches multiple Annex D.
        'BT-1252',
    }

    if 'DEBUG' in os.environ:
        (sourcedir / 'xlsx').mkdir(exist_ok=True)

    dfs = []
    for name, xlsx in excel_files():
        for sheet_name in xlsx.sheet_names:
            sheet = sheet_name

            # XXX: Typo in sheet name. (Corresponds to "SF22 vs eForm13" in the same workbook.)
            if sheet == 'eForm3 vs SF22':
                sheet = 'eForm13 vs SF22'

            if sheet in ignore_sheets or any(regex.search(sheet) for regex in ignore_sheet_regex):
                continue

            match = keep_sheet_regex.search(sheet)
            if not match:
                raise click.ClickException(f"The sheet name {sheet!r} doesn't match a known pattern.")

            form_2019 = match.group(1)
            form_2015 = match.group(2)

            # Read the Excel file. The first row is a title for the table. Per the "Legend" sheet, "---" means
            # "Outdent level, not considered (e.g. title of business groups)".
            df = pd.read_excel(xlsx, sheet_name, skiprows=[0], na_values='---')

            # Avoid empty or duplicate headings.
            df.rename(columns={
                df.columns[0]: 'Empty',  # was "" ("Unnamed: 0" after `read_excel`)
                df.columns[1]: 'Indent level',  # was "Level" or "" ("Unnamed: 1" after `read_excel`)
                df.columns[8]: 'Level',  # "Level.1" after `read_excel` if column 1 was "Level"
            }, errors='raise', inplace=True)

            # Remove rows with an empty "Level", for which we have no mapping information.
            df = df[df['Level'].notna()]
            if df.empty:  # eN40 vs F20
                continue

            # Remove the first column if it is empty.
            if df['Empty'].isna().all():
                df.drop(columns='Empty', inplace=True)
            else:
                raise click.ClickException('The first column was expected to be empty.')

            # Add notice number columns, using '1' instead of '01' to ease joins.
            df['2019 form'] = [[number.lstrip('0') for number in form_2019.split(',')] for i in df.index]
            df['2015 form'] = form_2015.lstrip('0')

            # Trim whitespace.
            df['Name'] = df['Name'].str.strip()
            # Normalize indices ("D1 - 1.1.1" to "D1.1.1.1").
            df['Level'] = df['Level'].str.replace(' - ', '.')

            # Make values easier to work with (must occur after `isna` above).
            df.fillna('', inplace=True)

            basename = f'eForm{form_2019}-F{form_2015}'
            if 'DEBUG' in os.environ:
                df.to_csv(sourcedir / 'xlsx' / f'{basename}.csv', index=False)

            # We don't have guidance for defence forms, and they don't use the same form indices in any case. As such,
            # we have no rows for forms 9, 6 (F16), 18 (F17), 31 (F18), 22 (F19), but we do for 27 (F15) and 3 (F08).
            if form_2015 in ('16', '17', '18', '19'):
                continue

            # `explode()` requires lists with the same number of elements, but an "Element" is not repeated if it is
            # the same for all "Level". Extra newlines also complicate things.
            for label, row in df.iterrows():
                bt_id = row['ID']
                levels = row['Level']
                elements = row['Element']
                location = f'{basename}: {bt_id}: '

                # Remove spurious newlines in "Element" values.
                elements = elements.replace('\n\n', '\n')
                if replace_newlines.search(elements):
                    # Display the newlines that were removed, for user to review.
                    highlight = replace_newlines.sub(lambda m: click.style(m.group(0), blink=True), elements)
                    # Replace outside f-string ("SyntaxError: f-string expression part cannot include a backslash").
                    highlight = highlight.replace('\n', '\\n')

                    click.echo(f'{location}removed \\n: {highlight}')
                    elements = replace_newlines.sub(' ', elements)

                # XXX: Hardcode corrections or cases requiring human interpretation.
                if (
                    form_2019 == '15'
                    and form_2015 == '07'
                    and bt_id == 'BT-18'
                    and levels == 'I.3.4.1.1'
                    and len(elements.split('\n')) == 2
                ):
                    # Add the missing "Level".
                    levels = 'I.3.4.1.1\nI.3.4.1.2'
                elif (
                    form_2019 == '18'
                    and form_2015 == '17'
                    and bt_id == 'BT-750'
                    and levels == 'III.2.1.1.1\nIII.2.2.2\nIII.2.2.3\nIII.2.3.1.1\nIII.2.3.1.2'
                    and len(elements.split('\n')) == 2
                ):
                    # Unambiguously repeat the "Element".
                    elements = dedent("""\
                    Information and formalities necessary for evaluating if the requirements are met:
                    Information and formalities necessary for evaluating if the requirements are met:
                    Minimum level(s) of standards possibly required: (if applicable)
                    Information and formalities necessary for evaluating if the requirements are met:
                    Minimum level(s) of standards possibly required: (if applicable)
                    """)

                if '\n' not in levels:
                    # Warn about remaining newlines (BT-531 in F17 and F18).
                    if '\n' in elements:
                        click.secho(f'{location}unexpected \\n: {repr(row["Element"])}', fg='yellow')

                    # Update the data frame before `continue`.
                    df.at[label, 'Level'] = levels
                    df.at[label, 'Element'] = elements

                    continue

                # Split values, after removing leading and trailing newlines (occurs in "Level" values).
                levels = [value.strip() for value in levels.strip("\n").split("\n")]
                elements = [value.strip() for value in elements.strip("\n").split("\n")]

                # Repeat "Element" as many times as there are "Level".
                length = len(levels)
                if length > len(elements):
                    elements *= length
                if length != len(elements):
                    click.secho(f'{location}size differs: {row["Level"]} {row["Element"]}', fg='red')

                # Remove rows with a B... "Level" after split, for which we have no mapping information.
                for i, (level, element) in enumerate(zip(levels, elements)):
                    if 'B' in level:  # "B.1", "B 4.2", "(Annex B)"
                        del levels[i]
                        del elements[i]

                # XXX: Correct the source file's incorrect mapping between 2019 BTs and 2015 indices.
                if form_2015 in ('23', '25'):
                    remove = remove_concession
                else:
                    remove = remove_standard
                if bt_id in remove:
                    i = levels.index(remove[bt_id])
                    levels.pop(i)
                    elements.pop(i)

                # Check for errors in the source file.
                sections = set(level.split('.', 1)[0] for level in levels)
                if len(sections) > 1 and bt_id not in ignore:
                    raise click.ClickException(
                        f'{location}{row["Name"]} cannot map to multiple sections ({", ".join(row["Level"])}). '
                        'Edit manage.py to correct the source file.'
                    )

                # Update the data frame.
                df.at[label, 'Level'] = levels
                df.at[label, 'Element'] = elements

            try:
                df = df.explode(['Level', 'Element'])
            except ValueError as e:
                raise click.ClickException(f'{sheet}: {e}')

            df = df.explode('2019 form')

            dfs.append(df)

    pd.concat(dfs, ignore_index=True).to_csv(eformsdir / 'bt-indices-mapping.csv', index=False)


@cli.command()
def update_ted_xml_indices():
    # TODO
    pass


@cli.command()
def extract_2015_guidance():
    """
    Concatenate guidance for the 2015 regulation.

    \b
    Create or update output/mapping/eForms/2015-guidance.csv:

    - Concatenate the CSV files for the 2015 regulation.
    - Merge the ted-xml-indices.csv file, which replaces the "index" column.
    """
    ignore = {
        # II.1.2 is moved from /OBJECT_CONTRACT/CPV_MAIN to /OBJECT_CONTRACT/CPV_MAIN/CPV_CODE.
        '/OBJECT_CONTRACT/CPV_MAIN': 'no index',
        # IV.2.4 is moved from /PROCEDURE/LANGUAGES to /PROCEDURE/LANGUAGES/LANGUAGE.
        '/PROCEDURE/LANGUAGES': 'no index',
        # The winner's address is at index V.2.3 on all forms except F13 (V.3.3). eForms collapses the two.
        '/RESULTS/AWARDED_PRIZE/WINNERS/WINNER/ADDRESS_WINNER': 'V.2.3',
        # This XPath is at index D1.1 (regular) or D2.1 (utilities). Anyhow, the eForms mapping doesn't use this index.
        '/PROCEDURE/PT_AWARD_CONTRACT_WITHOUT_CALL/D_ACCORDANCE_ARTICLE': 'D1.1/D2.1',
    }

    df_indices = pd.read_csv(eformsdir / 'ted-xml-indices.csv')

    dfs = []
    for path in sorted(mappingdir.glob('*.csv')):
        df = pd.read_csv(path)

        # Add the "index" column from the indices file.
        df = df.merge(df_indices, how='left', on='xpath', suffixes=('_old', ''))
        # Add a "file" column for the source of the row.
        df['file'] = path.stem.split('_')[0]

        # Check that the indices are consistent across files.
        for label, row in df.iterrows():
            xpath = row['xpath']
            x = row['index_old']
            y = row['index']
            if pd.notna(x) and x != y and not y == ignore.get(xpath):
                raise click.ClickException(f'{xpath} is {x} in {path.name} but {y} in ted-xml-indices.csv')

        # Drop the "index" column from the guidance file.
        df.drop(columns='index_old', inplace=True)

        dfs.append(df)

    # ignore_index is required, as each data frame repeats indices.
    pd.concat(dfs, ignore_index=True).to_csv(eformsdir / '2015-guidance.csv', index=False)


@cli.command()
@click.argument('filename', type=click.Path())
def update_with_2015_guidance():
    """
    Update FILE with eForms XPaths and 2015 guidance.

    \b
    Create or update FILE from bt-indices-mapping.csv, bt-xpath-mapping.csv and 2015-guidance.csv.
    """

    def add(data, current_row):
        current_row['XPATH'] = tuple(sorted(current_row['XPATH']))  # for briefer diff
        data.append(current_row)

    # Start with the eForms file that contains indices used by the 2015 guidance.
    df = pd.read_csv(eformsdir / 'bt-indices-mapping.csv')

    # Merge in the eForms XPaths. Sort by "BT ID" to simplify the for-loop's logic below.
    df_xpath = pd.read_csv(eformsdir / 'bt-xpath-mapping.csv').sort_values('BT ID')

    data = []

    current_row = {'BT ID': None, 'XPATH': []}

    for label, row in df_xpath.iterrows():
        if row['BT ID'] != current_row['BT ID']:
            if current_row['XPATH']:
                add(data, current_row)
            df_xpath.at[label, 'XPATH'] = [row['XPATH']]
            current_row = row
        else:
            current_row['XPATH'].append(row['XPATH'])

    add(data, current_row)

    df = df.merge(pd.DataFrame(data, columns=df_xpath.columns), how='left', left_on='ID', right_on='BT ID')

    # TODO: This sort changes the output!
    df.sort_values(['ID', '2019 form', '2015 form', 'Level', 'XPATH'], inplace=True)

    # Merge in 2015 guidance.
    df = df.merge(pd.read_csv(eformsdir / '2015-guidance.csv'), how='left', left_on='Level', right_on='index')

    # TODO: Do we need to add 'Level' (and maybe '2015 form'?), to allow the mapping to be context dependent?
    # Need to merge Level, guidance (anything else?) into an array.
    df.drop_duplicates(['2019 form', 'ID'], inplace=True)

    # Add two manually-edited columns.
    df.loc[df['guidance'].notna(), 'status'] = 'imported from standard forms'
    df['comments'] = ''

    df.sort_values(['2019 form', '2015 form', 'ID'], inplace=True)

    df.drop(columns=[
        # bt-indices-mapping.csv: Defer these columns to the 2019 regulation's annex.
        'Indent level',
        'Name',
        'Data type',
        'Repeatable',
        'Description',
        'Legal Status',
        # bt-xpath-mapping.csv: Defer these columns to the 2019 regulation's annex.
        'BT ID',  # merge column
        'BT Name',
        # 2015-guidance.csv: Only want guidance related to 2015 regulation.
        'xpath',
        'label-key',  # TODO: Lookup and add the English label?
        'index',  # merge column
        'file',
    ], inplace=True)

    # TODO: Remove this line once satisfied with comparison.
    df = df[['ID', 'Level', '2019 form', '2015 form', 'XPATH', 'guidance', 'status', 'comments']]

    df.to_json(eformsdir / f'eforms-guidance-pre.json', orient='records', indent=2)


@cli.command()
@click.argument('filename', type=click.Path(exists=True))
def update_with_annex(filename):
    """
    Update FILE with details from the 2019 regulation's annex.

    \b
    Create or update FILE from source/CELEX_32019R1780_EN_ANNEX_TABLE2_Extended.xlsx
    """
    # TODO: Create rows for all BT/notice pairs. Also, there shouldn't be any empty legal statuses... A problem in the index mapping source file?
    line = []
    previous_level = 0

    with pd.ExcelFile(sourcedir / 'CELEX_32019R1780_EN_ANNEX_TABLE2_Extended.xlsx') as xlsx:
        # A warning is issued, because the Excel file has an unsupported extension.
        df_annex = pd.read_excel(xlsx, 'Annex')

        # Use the notice number column names.
        check(df_annex.shape[1], 52, 'columns')
        df_annex.rename(columns={
            # 0:Level, 1:ID, 2:Name, 3:Data type, 4:Repeatable, 5:Description, 52:Fields not included in the legal text
            df_annex.columns[i]: df_annex.iloc[0][i] for i in range(6, 51)
        }, errors='raise', inplace=True)

        # Remove extra header rows.
        check(df_annex['ID'].isna().sum(), 3, 'extra header rows')
        df_annex = df_annex[df_annex['ID'].notna()]

        # Ensure there are no duplicates.
        df_annex.set_index('ID', verify_integrity=True, inplace=True)

        # Trim whitespace.
        df_annex['Name'] = df_annex['Name'].str.strip()

        # Add groups, to assist in mapping.
        df_annex['Business groups'] = pd.Series(dtype=object)
        for label, row in df_annex.iterrows():
            current_level = len(row['Level'])

            # Adjust the size of this line of the "tree", then update the head.
            if current_level > previous_level:
                line.append((None, None))
            elif current_level < previous_level:
                line = line[:current_level]
            line[-1] = [label, row['Name']]

            df_annex.at[label, 'Business groups'] = dict(line[:-1])
            previous_level = current_level

    df_guidance = pd.read_json(filename, orient='records')
    df_guidance['Business groups'] = pd.Series(dtype=object)

    for label, row in df_guidance.iterrows():
        annex_label = row['ID']

        for column in ('Business groups', 'Name', 'Data type', 'Repeatable', 'Description'):
            df_guidance.at[label, column] = df_annex.at[annex_label, column]

        status = df_annex.at[annex_label, row['2019 form']]
        if status in ('CM', 'EM'):  # CM: if "conditions" met, EM: if it "exists"
            status = 'M'
        df_guidance.at[label, 'Legal status'] = status

    for column in ('Legal status', 'guidance', 'status', 'comments'):
        df_guidance[column].fillna('', inplace=True)

    # TODO: Uncomment columns once added via other command.
    df_guidance = df_guidance[[
        # Identification
        'ID',
        '2019 form',
        '2015 form',

        # Manual columns
        'guidance',
        'status',
        'comments',

        # 2019 regulation's annex (excludes "Level", like "+++")
        'Business groups',
        'Name',
        'Data type',
        'Repeatable',
        'Description',
        'Legal status',

        # XPath mapping
        'XPATH',
        # 'Additional information',

        # Index mapping
        'Level',
        # 'Element',
    ]]

    df_guidance.to_json(filename, orient='records', indent=2)


@cli.command()
def statistics():
    """
    Print statistics on the progress of the guidance for the 2019 regulation.
    """
    df = pd.read_json(eformsdir / 'eforms-guidance.json', orient='records')

    df_terms = df.drop_duplicates(subset='ID')
    total_terms = df_terms.shape[0]
    done_terms = df_terms[df_terms['status'].str.startswith('done')].shape[0]

    total = df.shape[0]
    imported = df[df['status'] == 'imported from standard forms'].shape[0]
    done = df[df['status'].str.startswith('done')].shape[0]
    ready = imported + done
    no_issue_no_guidance = df[(df['status'] == '') & (df['guidance'] == '')].shape[0]

    df_issue = df[df['status'].str.startswith('issue')]
    issue = df_issue.shape[0]
    issue_no_guidance = df_issue[df_issue['guidance'] == ''].shape[0]

    legal_status = {}
    for value in ('M', 'O', ''):
        df_legal_status = df[df['Legal status'] == value]
        legal_status[value] = {
            'total': df_legal_status.shape[0],
            'done': df_legal_status[df_legal_status['status'].str.startswith('done')].shape[0],
            'issue': df_legal_status[df_legal_status['status'].str.startswith('issue')].shape[0],
        }

    click.echo(dedent(f"""\
    - BTs ready for review: {done_terms}/{total_terms} ({done_terms / total_terms:.1%})
    - Rows ready for review: {ready}/{total} ({ready / total:.1%})
        - Imported from 2015 guidance: {imported} ({imported / total:.1%})
        - Added or edited after import: {done} ({done / total:.1%})
        - Per legal status:
            - Mandatory: {legal_status['M']['done']}/{legal_status['M']['total']} ({legal_status['M']['done'] / legal_status['M']['total']:.1%}), {legal_status['M']['issue']} with open issues ({legal_status['M']['issue'] / legal_status['M']['total']:.1%})
            - Optional: {legal_status['O']['done']}/{legal_status['O']['total']} ({legal_status['O']['done'] / legal_status['O']['total']:.1%}), {legal_status['O']['issue']} with open issues ({legal_status['O']['issue'] / legal_status['O']['total']:.1%})
            - Empty: {legal_status['']['done']}/{legal_status['']['total']} ({legal_status['']['done'] / legal_status['']['total']:.1%}), {legal_status['']['issue']} with open issues ({legal_status['']['issue'] / legal_status['']['total']:.1%})
    - Rows with [open issues](https://github.com/open-contracting/european-union-support/labels/eforms): {issue} ({issue / total:.1%}), {issue_no_guidance} without guidance
    - Rows without issues and without guidance: {no_issue_no_guidance} ({no_issue_no_guidance / total:.1%})
    """))  # noqa: E501


@cli.command()
@click.argument('file', type=click.File())
def fields_without_extensions(file):
    """
    Print undefined fields in the guidance for the 2015 regulation.
    """
    subjects = {
        # Unambiguous
        'award': 'awards',
        'contract': 'contracts',
        'lot': 'tender/lots',
        'party': 'parties',
        'release': '',
        'statistic': 'bids/statistics',
        'charge': 'contracts/implementation/charges',

        # Ambiguous
        'amendment': {
            'CHANGES': 'tender/amendments',
            'MODIFICATIONS_CONTRACT': 'contracts/amendments',
        },
        'classification': {
            'CONTRACTING_BODY': 'parties/details/classifications',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
        },
        'criterion': {
            'LEFTI': 'tender/selectionCriteria/criteria',
            'OBJECT_CONTRACT': 'tender/lots/awardCriteria/criteria',
        },
        'item': {
            'MODIFICATIONS_CONTRACT': 'contracts/items',
            'OBJECT_CONTRACT': 'tender/items',
        },
        'object': {
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
        },
    }

    unknowns = {
        # Unambiguous
        '.additionalContactPoints': 'parties',
        '.awardCriteria': 'tender',
        '.awardID': 'contracts',
        '.countryCode': 'parties/address',
        '.details.classifications': 'parties',
        '.documentType': 'tender/documents',
        '.estimatedValue.amount': 'contracts/implementation/charges',
        '.financingParty.id': 'planning/budget/finance',
        '.financingParty.name': 'planning/budget/finance',
        '.identifier.id': 'parties',
        '.identifier.legalName': 'parties',
        '.identifier.scheme': 'parties',
        '.measure': 'bids/statistics',  # metrics extension not used
        '.minimum': 'tender/selectionCriteria/criteria',
        '.paidBy': 'contracts/implementation/charges',
        '.roles': 'parties',
        '.secondStage.maximumCandidates': 'tender/lots',
        '.secondStage.minimumCandidates': 'tender/lots',
        '.subcontracting.maximumPercentage': 'awards',
        '.suppliers': 'awards',  # contract suppliers extension not used

        # Ambiguous
        '.additionalClassifications': {
            'MODIFICATIONS_CONTRACT': 'contracts/items',
        },
        '.description': {
            # Root
            'LEFTI': 'tender/selectionCriteria/criteria',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
            # XPath
            '/CONTRACTING_BODY/CA_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY_OTHER': 'parties/details/classifications',
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CATEGORY/@CTYPE': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
        },
        '.id': {
            # Root
            'CHANGES': 'tender/amendments',
            'LEFTI': 'tender/documents',
            'PROCEDURE': 'tender/procurementMethodRationaleClassifications',
            # XPath
            '/AWARD_CONTRACT': 'awards',
            '/AWARD_CONTRACT/AWARDED_CONTRACT': 'contracts',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/CONTRACTORS/CONTRACTOR/ADDRESS_CONTRACTOR': 'awards/suppliers',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/CONTRACTORS/CONTRACTOR/ADDRESS_PARTY': 'parties/shareholders',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_EMEANS': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_NON_EU': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_OTHER_EU': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/TENDERS/NB_TENDERS_RECEIVED_SME': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VAL_PRICE_PAYMENT': 'contracts/implementation/charges',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VAL_REVENUE': 'contracts/implementation/charges',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VALUES/VAL_RANGE_TOTAL/HIGH': 'bids/statistics',
            '/AWARD_CONTRACT/AWARDED_CONTRACT/VALUES/VAL_RANGE_TOTAL/LOW': 'bids/statistics',
            '/CONTRACTING_BODY/ADDRESS_CONTRACTING_BODY': 'parties',
            '/CONTRACTING_BODY/DOCUMENT_RESTRICTED': 'tender/participationFees',
            '/CONTRACTING_BODY/CA_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_ACTIVITY_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CA_TYPE_OTHER': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY/@VALUE': 'parties/details/classifications',
            '/CONTRACTING_BODY/CE_ACTIVITY_OTHER': 'parties/details/classifications',
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CONTRACTORS/CONTRACTOR/ADDRESS_CONTRACTOR': 'awards/suppliers',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_ADDITIONAL/CPV_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'contracts/items/additionalClassifications',  # noqa: E501
            '/MODIFICATIONS_CONTRACT/INFO_MODIFICATIONS': 'contracts/amendments',
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CATEGORY/@CTYPE': 'tender/additionalClassifications',
            '/OBJECT_CONTRACT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',  # noqa: E501
            '/OBJECT_CONTRACT/OBJECT_DESCR/EU_PROGR_RELATED': 'planning/budget/finance',
            '/OBJECT_CONTRACT/VAL_RANGE_TOTAL/HIGH': 'bids/statistics',
            '/OBJECT_CONTRACT/VAL_RANGE_TOTAL/LOW': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE': 'contracts',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS_OTHER_EU': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/PARTICIPANTS/NB_PARTICIPANTS_SME': 'bids/statistics',
            '/RESULTS/AWARDED_PRIZE/WINNERS/WINNER/ADDRESS_WINNER': 'awards/suppliers',
        },
        '.items': {
            'MODIFICATIONS_CONTRACT': 'contracts',
        },
        '.name': {
            # Root
            'AWARD_CONTRACT': 'awards/suppliers',
            'MODIFICATIONS_CONTRACT': 'awards/suppliers',
            'OBJECT_CONTRACT': 'parties',
            'RESULTS': 'awards/suppliers',
            # XPath
            '/CONTRACTING_BODY/ADDRESS_CONTRACTING_BODY': 'buyer',
            '/PROCEDURE/PARTICIPANT_NAME': 'parties',
            '/PROCEDURE/MEMBER_NAME': 'tender/designContest/juryMembers',
        },
        '.newValue': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.classifications': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.date': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.newValue.text': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.classifications': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.date': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.oldValue.text': {
            'CHANGES': 'tender/amendments/unstructuredChanges'
        },
        '.region': {
            '/OBJECT_CONTRACT/OBJECT_DESCR/NUTS': 'tender/items/deliveryAddresses',
            '/MODIFICATIONS_CONTRACT/DESCRIPTION_PROCUREMENT/NUTS': 'contracts/items/deliveryAddresses',
        },
        '.relatedLot': {
            'AWARD_CONTRACT': 'bids/statistics',
            'CHANGES': 'tender/amendments/unstructuredChanges',
            'RESULTS': 'bids/statistics',
        },
        '.relatedLots': {
            'AWARD_CONTRACT': 'awards',
            'OBJECT_CONTRACT': 'planning/budget/finance',
        },
        '.scheme': {
            'MODIFICATIONS_CONTRACT': 'contracts/items/additionalClassifications',
            # XPath
            '/CHANGES/CHANGE/NEW_VALUE/CPV_ADDITIONAL/CPV_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_MAIN/CPV_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/NEW_VALUE/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/newValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_ADDITIONAL/CPV_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_MAIN/CPV_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/CHANGES/CHANGE/OLD_VALUE/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/amendments/unstructuredChanges/oldValue/classifications',  # noqa: E501
            '/OBJECT_CONTRACT/CPV_MAIN/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_CODE': 'tender/items/additionalClassifications',
            '/OBJECT_CONTRACT/OBJECT_DESCR/CPV_ADDITIONAL/CPV_SUPPLEMENTARY_CODE': 'tender/items/additionalClassifications',  # noqa: E501
            '/OBJECT_CONTRACT/CATEGORY': 'tender/additionalClassifications',
        },
        '.shareholder.id': {
            'AWARD_CONTRACT': 'parties/shareholders',
        },
        '.shareholder.name': {
            'AWARD_CONTRACT': 'parties/shareholders',
        },
        '.status': {
            'AWARD_CONTRACT': 'contracts',
            'RESULTS': 'contracts',
        },
        '.title': {
            'AWARD_CONTRACT': 'contracts',
            'LEFTI': 'tender/targets',
        },
        '.type': {
            # Root
            'CONTRACTING_BODY': 'tender/participationFees',
            'LEFTI': 'tender.selectionCriteria.criteria',
            'OBJECT_CONTRACT': 'tender/lots/awardCriteria/criteria',
        },
        '.value': {
            'AWARD_CONTRACT': 'bids/statistics',
            'OBJECT_CONTRACT': 'bids/statistics',
            'RESULTS': 'bids/statistics',
        },
        '.where': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
        '.where.label': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
        '.where.section': {
            'CHANGES': 'tender/amendments/unstructuredChanges',
        },
    }

    unhandled = set()
    paths = set()

    def report(path, row):
        value = (path, row.get('xpath'))
        if value not in unhandled:
            print(f"unhandled: {path} ({row.get('xpath')}: {row['guidance']})", err=True)
        unhandled.add(value)

    for path in (mappingdir, mappingdir / 'shared'):
        for filename in path.glob('*.csv'):
            with filename.open() as f:
                reader = csv.DictReader(f)
                for row in reader:
                    if row.get('guidance'):
                        for match in re.finditer(r"(?:([a-z]+)'s )?\[?`([^`]+)`", row['guidance']):
                            path = match.group(2)
                            if path in ('true', 'false', 'value'):  # JSON boolean, exceptional case
                                continue
                            if re.search(r'^[A-Z][a-z][A-Za-z]+$', path):  # JSON Schema definition
                                continue
                            if re.search(r'^(/[A-Z_]+)+$', path):  # XPath
                                continue
                            if re.search(r'^[A-Z_]+$', path):  # XML element
                                continue

                            subject = match.group(1)

                            prefix = ''
                            if subject:
                                try:
                                    prefix = subjects[subject]
                                except KeyError as e:
                                    click.echo(f"KeyError: Add a {e} key to the `subjects` list")
                            elif path in unknowns:
                                try:
                                    prefix = unknowns[path]
                                except KeyError as e:
                                    click.echo(f"KeyError: Add a {e} key to the `unknowns` list")
                            elif path[0] == '.':
                                report(path, row)
                                continue

                            if isinstance(prefix, dict):
                                xpath = row.get('xpath', '/')
                                root = xpath.split('/', 2)[1]
                                if root in prefix:
                                    key = root
                                else:
                                    key = xpath
                                try:
                                    prefix = prefix[key]
                                except KeyError:
                                    report(path, row)
                                    continue

                            path = prefix + path
                            paths.add(path.replace('.', '/'))

    seen = [row['path'] for row in csv.DictReader(file)]
    for path in sorted(list(paths)):
        if path not in seen:
            click.echo(path)

    # Uncomment to print all the paths for a specific object.
    # for path in sorted(list(paths)):
    #     if '/items' in path:
    #         click.echo(path)


if __name__ == '__main__':
    cli()
